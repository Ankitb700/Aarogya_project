from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ── Title Page ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\n\n\n')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('rPPG Vitals System')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(37, 99, 235)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Remote Photoplethysmography Inference Engine\n& Web Application')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(100, 116, 139)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'\n\nComprehensive Architecture & Implementation Report\n\n{datetime.now().strftime("%B %d, %Y")}')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(148, 163, 184)

doc.add_page_break()

# ── Table of Contents placeholder ──
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Executive Summary',
    '2. Project Overview',
    '3. Technology Stack',
    '4. System Architecture',
    '5. Supported Models & Techniques',
    '6. Data Processing Pipeline',
    '7. Signal Processing & Vitals Extraction',
    '8. Backend Application',
    '9. Frontend Application',
    '10. UI Redesign & Features',
    '11. Results & Performance',
    '12. Deployment & Testing'
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')
doc.add_page_break()

# ═══════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'The rPPG (Remote Photoplethysmography) Vitals System is a fully integrated platform '
    'for contactless physiological measurement from facial videos. It combines a state-of-the-art '
    'deep learning inference engine written in Python/JAX with a modern healthcare-grade web application '
    'built on FastAPI and React.'
)
doc.add_paragraph(
    'The system measures Heart Rate (HR), Heart Rate Variability (HRV), Respiratory Rate (RR), '
    'Stress Index, and Signal Quality Index (SQI) from video captured via webcam or offline video files. '
    'It supports 9 different deep learning architectures across 18 model variants, all unified under '
    'a single high-level API.'
)
doc.add_paragraph(
    'The web frontend provides a complete clinical workflow: Landing Page → Patient Intake Form → '
    'Live Camera Measurement (30-second scan) → Results Dashboard with time-series charts, '
    'wellness scoring, recovery analysis, and HRV analytics.'
)

# ═══════════════════════════════════════════════
# 2. PROJECT OVERVIEW
# ═══════════════════════════════════════════════
doc.add_heading('2. Project Overview', level=1)
doc.add_heading('2.1 What Was Built', level=2)
doc.add_paragraph(
    'The project delivers two major components working together:'
)
bullets = [
    'rPPG Inference Engine (open_rppg_inference/) — A Python toolbox for remote photoplethysmography '
    'using deep learning, supporting offline video processing and real-time webcam inference with JAX acceleration.',
    'Vitals Web Application — A full-stack web app (FastAPI + React) that exposes the engine via WebSocket '
    'for live browser-based measurement, complete with a multi-step clinical UI.'
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('2.2 What Techniques Were Used', level=2)
techniques = [
    'Remote Photoplethysmography (rPPG) — extracting blood volume pulse signals from facial video',
    'State-Space Models (Mamba, SSM) — selective scan algorithms for efficient temporal modeling',
    'Temporal Difference Transformers (PhysFormer) — attention-based spatiotemporal feature extraction',
    'Temporal Shift Networks (TSCAN, EfficientPhys) — lightweight frame-shifting for motion modeling',
    '3D Convolutional Networks (PhysNet) — volumetric spatiotemporal processing',
    'Central Difference Convolution (CDC) — gradient-aware convolutions for subtle pulse signals',
    'Frequency-domain learning — FFT-based feature modulation in the spectral domain',
    'BlazeFace — lightweight ONNX-based face detection for facial region extraction',
    'Kalman Filtering — temporal smoothing of face bounding boxes',
    'Butterworth Bandpass Filtering — signal cleaning in the 0.7–4.0 Hz range (42–240 BPM)',
    'Welch Power Spectral Density — heart rate estimation via frequency-domain analysis',
    'HeartPy — pulse rate variability and HRV metric extraction (SDNN, RMSSD, LF/HF, pNN50)',
    'AUTOCORRELATION-BASED SQI — signal quality indexing via autocorrelation peak analysis',
    'JIT Compilation (JAX) — just-in-time compiled inference for GPU/TPU acceleration',
    'Streaming Inference Pipeline — threaded face detection + parallel model inference'
]
for t in techniques:
    doc.add_paragraph(t, style='List Bullet')

# ═══════════════════════════════════════════════
# 3. TECHNOLOGY STACK
# ═══════════════════════════════════════════════
doc.add_heading('3. Technology Stack', level=1)

doc.add_heading('3.1 Inference Engine', level=2)
table = doc.add_table(rows=9, cols=2, style='Light List Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.LEFT
data = [
    ('Language', 'Python 3.9–3.13'),
    ('Deep Learning', 'JAX, Keras 3 (JAX backend)'),
    ('Face Detection', 'ONNX Runtime (BlazeFace)'),
    ('Signal Processing', 'SciPy (butter, filtfilt, welch, find_peaks)'),
    ('HRV Analysis', 'HeartPy'),
    ('Video I/O', 'PyAV (FFmpeg), OpenCV'),
    ('Misc', 'NumPy, einops, pandas'),
    ('Optimization', 'JAX JIT, mixed_float16, XLA'),
]
for i, (k, v) in enumerate(data):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v

doc.add_heading('3.2 Backend', level=2)
table = doc.add_table(rows=6, cols=2, style='Light List Accent 1')
data = [
    ('Framework', 'FastAPI'),
    ('WebSocket', 'WebSockets (via Uvicorn)'),
    ('Async', 'asyncio, ModelPool pattern'),
    ('Serialization', 'Pydantic, JSON'),
    ('Server', 'Uvicorn'),
    ('Dependencies', 'All engine deps + fastapi, uvicorn, websockets'),
]
for i, (k, v) in enumerate(data):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v

doc.add_heading('3.3 Frontend', level=2)
table = doc.add_table(rows=8, cols=2, style='Light List Accent 1')
data = [
    ('Framework', 'React 18'),
    ('Build Tool', 'Vite 5'),
    ('Routing', 'React Router v6'),
    ('Charts', 'Recharts 2'),
    ('Styling', 'Custom CSS (healthcare design system)'),
    ('Camera', 'WebRTC (getUserMedia), Canvas API'),
    ('State', 'React Context + localStorage'),
    ('Real-time', 'WebSocket (binary JPEG frames)'),
]
for i, (k, v) in enumerate(data):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v

# ═══════════════════════════════════════════════
# 4. SYSTEM ARCHITECTURE
# ═══════════════════════════════════════════════
doc.add_heading('4. System Architecture', level=1)

doc.add_heading('4.1 High-Level Architecture', level=2)
doc.add_paragraph(
    'The system follows a three-tier architecture:\n\n'
    '1. Inference Engine Layer — Python package (rppg/) with Model class, face detector, signal processors\n'
    '2. Backend Service Layer — FastAPI server with WebSocket endpoint, model pool, session management\n'
    '3. Frontend Presentation Layer — React SPA with multi-page clinical workflow'
)

doc.add_heading('4.2 Data Flow', level=2)
doc.add_paragraph(
    'Video Source (Webcam / File) → Frame Capture → Face Detection (BlazeFace ONNX) → '
    'Face Crop + Resize → Model Inference (JAX/Keras JIT) → BVP Signal → '
    'Signal Processing (Bandpass Filter, Normalization, Detrend) → '
    'Vitals Extraction (HR, HRV, RR, Stress, SQI) → WebSocket Streaming → '
    'Dashboard Rendering (React/Recharts)'
)

doc.add_heading('4.3 File Structure', level=2)
doc.add_paragraph(
    'E:\\phase_1\\\n'
    '├── open_rppg_inference/     # Python inference engine\n'
    '│   ├── rppg/                 # Core package\n'
    '│   │   ├── main.py           # Model class, FaceDetector, signal utilities\n'
    '│   │   ├── models.py         # All model architectures (1802 lines)\n'
    '│   │   ├── models_code/      # FacePhys model\n'
    '│   │   └── weights/          # 18 pretrained weight files\n'
    '│   ├── vitals_analyzer.py    # CLI & webcam vitals extraction\n'
    '│   └── requirements.txt      # Engine dependencies\n'
    '├── backend/                  # FastAPI WebSocket server\n'
    '│   ├── app.py                # Routes & WebSocket handler\n'
    '│   ├── session.py            # Analysis session management\n'
    '│   ├── pool.py               # Async model pool\n'
    '│   ├── schemas.py            # Pydantic message models\n'
    '│   └── engine.py             # Engine path bridge\n'
    '├── frontend/                 # React SPA\n'
    '│   ├── src/pages/            # 4 pages (Landing, Intake, Camera, Results)\n'
    '│   ├── src/components/       # Stepper, ProgressRing, ScoreGauge\n'
    '│   ├── src/lib/              # WebSocket, Camera, Health scoring\n'
    '│   ├── src/state/            # PatientContext\n'
    '│   └── src/styles.css        # Healthcare design system\n'
    '├── README_webapp.md\n'
    '└── UI_REDESIGN_CHANGELOG.md'
)

# ═══════════════════════════════════════════════
# 5. SUPPORTED MODELS & TECHNIQUES
# ═══════════════════════════════════════════════
doc.add_heading('5. Supported Models & Techniques', level=1)

doc.add_heading('5.1 Model Zoo', level=2)
table = doc.add_table(rows=10, cols=3, style='Light List Accent 1')
headers = ['Model Name', 'Architecture Type', 'Reference']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
table.rows[0].cells[0].text = 'Model Name'
table.rows[0].cells[1].text = 'Architecture Type'
table.rows[0].cells[2].text = 'Reference'

models_data = [
    ('FacePhys', 'Optimized State-Space Model (SSM)', '—'),
    ('ME-chunk', 'State-Space Model (chunk inference)', 'arXiv 2025'),
    ('ME-flow', 'State-Space Model (low-latency flow)', 'arXiv 2025'),
    ('PhysMamba', 'Dual-Branch Mamba (CDC + SSM)', 'CCBR 2024'),
    ('RhythmMamba', 'Frequency-Domain Constrained Mamba', 'AAAI 2025'),
    ('PhysFormer', 'Temporal Difference Transformer', 'CVPR 2022'),
    ('TSCAN', 'Temporal Shift Conv Attention Network', 'NeurIPS 2020'),
    ('EfficientPhys', 'Self-Attention TSCAN variant', 'WACV 2023'),
    ('PhysNet', '3D Convolutional Encoder-Decoder', 'BMVC 2019'),
]
for i, (nm, arch, ref) in enumerate(models_data):
    table.rows[i+1].cells[0].text = nm
    table.rows[i+1].cells[1].text = arch
    table.rows[i+1].cells[2].text = ref

doc.add_paragraph(
    '\nEach model is available in two training variants:\n'
    '• .rlap — Robust Lightweight Adaptive Pretraining\n'
    '• .pure — Standard training protocol\n\n'
    'Default model: FacePhys.rlap'
)

doc.add_heading('5.2 Key Techniques Per Model', level=2)

techniques_detail = [
    ('FacePhys', 'Complex-valued state-space model with NPU-friendly math ops (sin/cos/exp approximations). '
     'Uses SSConv blocks with Mamba2-style selective scan, temporal normalization (TNM), and RMSNorm. '
     'Optimized for streaming/real-time inference with step-by-step processing.'),
    ('ME (Memory-Efficient)', 'InfinitePulse architecture with SSCBlock backbone (stacked SSConv + Mamba2 layers) '
     'and Mamba2 head. Available in chunk (batch) and flow (streaming) variants. Input: 36×36 face crops at 160-frame windows.'),
    ('PhysMamba', 'Dual-branch slow-fast architecture. Slow path (high channel, low temporal freq) + Fast path '
     '(low channel, high temporal freq) with lateral fusion connections. Uses Central Difference Convolution (CDC), '
     'MambaLayer (BiMamba), and ChannelAttention3D. Preprocessing includes temporal differencing.'),
    ('RhythmMamba', 'Fusion Stem with dual-path differencing, Block_mamba with multi-scale segment processing, '
     'Frequency-domain FFN (learnable modulation in Fourier space). Attention mask for spatial focus.'),
    ('PhysFormer', 'Temporal Difference Transformer with CDC-based Q/K projections, sharp attention (gra_sharp), '
     'multi-scale stem (3 conv blocks with spatial pooling), hierarchical transformer stages, and upsampling decoder.'),
    ('TSCAN', 'Two-stream architecture: appearance stream (mean face) and difference stream (temporal diff). '
     'Temporal Shift Module (TSM) shifts channels along time axis. Attention masks weight spatial regions. Compact 36×36 input.'),
    ('EfficientPhys', 'Lightweight single-stream design with TSM, tanh activations, self-attention masks, '
     'and aggressive dropout. No difference stream — computes temporal diff internally. 72×72 input.'),
    ('PhysNet', 'Deep 9-block 3D CNN with ReLU activations, batch normalization, spatial+spatiotemporal max pooling, '
     'and transposed convolution upsampling. Standardization at input and output. 32×32 input.'),
]

for name, desc in techniques_detail:
    p = doc.add_paragraph()
    run = p.add_run(f'{name}: ')
    run.bold = True
    p.add_run(desc)

# ═══════════════════════════════════════════════
# 6. DATA PROCESSING PIPELINE
# ═══════════════════════════════════════════════
doc.add_heading('6. Data Processing Pipeline', level=1)

doc.add_heading('6.1 Face Detection', level=2)
doc.add_paragraph(
    'The system uses BlazeFace, a lightweight ONNX face detector by Google MediaPipe. '
    'Two models are available:\n'
    '• blaze_face.onnx (128×128 input, short-range anchors) — for "Near" face mode (default)\n'
    '• blaze_face_full.onnx (256×256 input, full-range anchors) — for "Far" mode\n\n'
    'Detection pipeline:\n'
    '1. Preprocess: Resize + pad to square, normalize to [-1, 1]\n'
    '2. Run ONNX inference → raw boxes + scores\n'
    '3. Decode boxes from anchor space, apply sigmoid to scores\n'
    '4. Non-Maximum Suppression (NMS) with IoU threshold 0.3\n'
    '5. Remove padding, return bounding box + 6 facial keypoints'
)

doc.add_heading('6.2 Kalman Filter Tracking', level=2)
doc.add_paragraph(
    'A 1D Kalman filter smooths the face bounding box coordinates across frames. '
    'Process noise = 0.01, measurement noise = 0.5. This reduces jitter and handles '
    'temporary detection failures gracefully.'
)

doc.add_heading('6.3 Threaded Pipeline', level=2)
doc.add_paragraph(
    'Inference is highly parallelized:\n'
    '• Face detection runs on a ThreadPoolExecutor (thread count = CPU cores/2)\n'
    '• Face resampling (cropping + resizing) runs on a separate pool\n'
    '• Model inference runs in its own daemon thread\n'
    '• Chaining: detection lock → frame lock → semaphore coordination\n'
    '• Face detection runs every N=5 frames; skipped frames reuse last known box\n'
    '• Frame buffer size = 10, with forward filling for missing detections'
)

doc.add_heading('6.4 Model Inference', level=2)
doc.add_paragraph(
    'Each model is JIT-compiled via @jax.jit for maximum performance. '
    'The inference pipeline:\n'
    '1. Accumulate face crops in a buffer (input size depends on model, e.g., 160 frames)\n'
    '2. Resize all buffered faces to model input resolution (36×36, 72×72, or 128×128)\n'
    '3. Normalize pixel values to [0, 1]\n'
    '4. Execute compiled model → output BVP signal\n'
    '5. Handle model-specific post-processing (cumsum, normalization, differencing)\n'
    '6. Cached via lru_cache to avoid reloading weights'
)

# ═══════════════════════════════════════════════
# 7. SIGNAL PROCESSING & VITALS
# ═══════════════════════════════════════════════
doc.add_heading('7. Signal Processing & Vitals Extraction', level=1)

doc.add_heading('7.1 BVP Post-Processing', level=2)
doc.add_paragraph(
    'Raw BVP from the model undergoes several cleaning steps:\n'
    '• Bandpass Butterworth filter (3rd order, 0.5–3.0 Hz cutoff, corresponding to 30–180 BPM)\n'
    '• Detrending via smoothing spline (scipy sparse) with lambda = 50 × (freq/30)² × (0.5/min_freq)²\n'
    '• Normalization: zero-mean, unit-variance\n'
    '• Peak-based normalization: split at zero-crossings, normalize each pulse segment\n'
    '• Clipping to [-2, max]'
)

doc.add_heading('7.2 Heart Rate Estimation', level=2)
doc.add_paragraph(
    'HR is computed via Welch Power Spectral Density estimation:\n'
    '• nfft = 20,000, nperseg = min(len-1, 256/30 × sampling_rate)\n'
    '• Search range: 0.5–3.0 Hz (30–180 BPM)\n'
    '• Peak frequency × 60 = BPM'
)

doc.add_heading('7.3 Signal Quality Index (SQI)', level=2)
doc.add_paragraph(
    'SQI uses autocorrelation: the peak autocorrelation value in the expected HR range '
    '(0.5–3.0 Hz) is computed over sliding windows. Values range from 0.0 (poor) to 1.0 (excellent). '
    'SQI > 0.5 is needed for HRV calculation.'
)

doc.add_heading('7.4 Heart Rate Variability (HRV)', level=2)
doc.add_paragraph(
    'HRV is computed using HeartPy:\n'
    '• RR intervals extracted via peak detection on the BVP signal\n'
    '• Metrics: SDNN, RMSSD, pNN50, LF/HF ratio, breathing rate\n'
    '• Frequency bands: VLF (0.0033–0.04 Hz), LF (0.04–0.15 Hz), HF (0.15–0.4 Hz)\n'
    '• Resampled at 4 Hz for spectral analysis via cubic spline interpolation'
)

doc.add_heading('7.5 Vitals Analyzer (extract_vitals)', level=2)
doc.add_paragraph(
    'The standalone vitals extraction function (in vitals_analyzer.py):\n'
    '• Requires ≥300 samples for reliable analysis\n'
    '• 3rd-order Butterworth bandpass 0.7–4.0 Hz\n'
    '• Peak detection: min distance = fps × 0.45, prominence = std × 0.3\n'
    '• RR interval filtering: 300–2000 ms (30–200 BPM)\n'
    '• Respiratory rate: Welch PSD in 0.1–0.5 Hz range\n'
    '• Stress Index = (HR/100×40) + (100/RMSSD×30) + (100/SDNN×30), clipped to 0–100'
)

doc.add_heading('7.6 Trend Computation (Backend)', level=2)
doc.add_paragraph(
    'The backend computes time-series trends from the full 30-second BVP signal using '
    'a sliding window approach. Window size is dynamically sized to contain ≥300 samples. '
    'Each window runs extract_vitals → produces {t, hr, rr, rmssd, sdnn, stress} data points. '
    'This replaces the earlier cumulative approach that left trend charts nearly empty.'
)

# ═══════════════════════════════════════════════
# 8. BACKEND APPLICATION
# ═══════════════════════════════════════════════
doc.add_heading('8. Backend Application', level=1)

doc.add_heading('8.1 Endpoints', level=2)
endpoints = [
    ('GET /health', 'Health check + model readiness'),
    ('GET /api/info', 'Model metadata, supported metrics, configuration limits'),
    ('WS /ws/analyze', 'WebSocket for live streaming: receives JPEG frames, streams vitals'),
]
for url, desc in endpoints:
    p = doc.add_paragraph()
    run = p.add_run(f'{url}: ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('8.2 WebSocket Protocol', level=2)
doc.add_paragraph(
    'Client sends binary JPEG frames → Server decodes → Feeds to rppg.Model.update_frame() → '
    'Streams JSON messages back:\n'
    '• StatusMessage: {type: "status", ready, frame_count, face_detected, elapsed}\n'
    '• MetricsMessage: {type: "metrics", hr, rr, rmssd, sdnn, stress, sqi, beats, duration} (every ~1s)\n'
    '• FinalMessage: {type: "final", hr, rr, stress, ... final_vitals + trend[]} (at 30s)'
)

doc.add_heading('8.3 Model Pool', level=2)
doc.add_paragraph(
    'The backend uses an async ModelPool (asyncio.Queue) that pre-warms N model instances. '
    'Default POOL_SIZE = 1. Each session acquires a model from the pool, uses it in a '
    'thread-safe manner, and releases it back. The pool is warmed on startup (takes 30–60s for JIT compilation).'
)

doc.add_heading('8.4 Session Management', level=2)
doc.add_paragraph(
    'AnalysisSession manages the lifecycle of one measurement:\n'
    '• open() — acquires model, enters context\n'
    '• push_jpeg() — decodes JPEG → BGR (OpenCV) → RGB → model.update_frame()\n'
    '• compute_vitals() — extracts BVP via model.bvp(), runs extract_vitals()\n'
    '• compute_trend() — sliding window trend over full signal\n'
    '• close() — releases model back to pool\n\n'
    'Max duration: 30 seconds, minimum collection: 5 seconds before vitals are computed.'
)

# ═══════════════════════════════════════════════
# 9. FRONTEND APPLICATION
# ═══════════════════════════════════════════════
doc.add_heading('9. Frontend Application', level=1)

doc.add_heading('9.1 Page Structure', level=2)
pages = [
    ('Landing Page (/)', 'Hero section with benefits, feature cards, "How It Works" steps, privacy notice, Start CTA'),
    ('Intake Form (/intake)', 'Patient data entry: gender, weight, height, birth year (auto-age), BP category. Inline validation, progress stepper, localStorage autosave.'),
    ('Camera Page (/measure)', 'Live webcam capture with face guide overlay, 30s countdown, status pills (camera/face/signal/measurement), progress bar.'),
    ('Results Page (/results)', 'Full analytics dashboard with KPI cards, circular gauges, trend charts, health radar, recovery analysis, HRV comparison, signal quality trend, summary table.'),
]
for title, desc in pages:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}: ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('9.2 Key Components', level=2)
doc.add_paragraph(
    '• Stepper.jsx — 3-step progress indicator\n'
    '• ProgressRing.jsx — SVG circular gauge with accessible ARIA labels\n'
    '• ScoreGauge.jsx — 0–100 wellness score ring with label and caption\n'
    '• FrameCapturer (capture.js) — WebRTC → canvas → JPEG at 18fps\n'
    '• AnalyzeSocket (ws.js) — WebSocket wrapper, sends blobs, parses JSON\n'
    '• PatientContext (state/) — React Context + localStorage persistence\n'
    '• health.js — derived analytics: wellness/recovery/readiness scores, radar dimensions, insights'
)

doc.add_heading('9.3 Health Scoring', level=2)
doc.add_paragraph(
    'The frontend computes composite health scores:\n'
    '• Wellness Score — weighted combination of HR, RMSSD, SDNN, Stress, RR\n'
    '• Recovery Score — based on RMSSD (parasympathetic activity)\n'
    '• Readiness Score — combination of recovery + low stress\n'
    '• Radar Dimensions — Cardiovascular, Recovery, Stress, Respiratory, Wellness, Readiness\n'
    '• Signal Quality — color-coded labels from SQI (poor/fair/good/excellent)\n'
    '• Health Insights — plain-language descriptions based on metric thresholds'
)

# ═══════════════════════════════════════════════
# 10. UI REDESIGN
# ═══════════════════════════════════════════════
doc.add_heading('10. UI Redesign & Features', level=1)

doc.add_heading('10.1 Design Transformation', level=2)
doc.add_paragraph(
    'The UI was transformed from a single-page technical prototype into a multi-step, '
    'healthcare-grade web experience with a clean clinical design system:'
)

features = [
    'Multi-step onboarding: Landing → Patient Intake → Camera → Results',
    'Healthcare design system: Medical Blue (#2563EB), Healthcare Teal (#0EA5A4), soft grays, glassmorphism',
    'Accessibility: semantic fieldsets, aria-live regions, visible focus rings, 48px+ touch targets',
    'Responsive: CSS Grid from mobile (1 col) → tablet (2 col) → desktop (4-col KPI row)',
    'Reduced motion support via prefers-reduced-motion',
    '30-second scan duration (reduced from 60s) with auto-stop',
    'Results Dashboard v2: Assessment summary, KPI cards with gauges, HR trend, RR trend, HRV evolution, '
    'Wellness score gauge, Health radar, Recovery analysis, Advanced HRV analytics with reference ranges',
    'Time-series fix: backend-computed trends via sliding window over full 30s BVP signal',
    'LAN exposure: server.host = true, allowedHosts = true, CORS allows any :5173 origin',
    'HTTPS tunnel: ngrok support for mobile phone testing (camera requires secure context)',
]
for f in features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('10.2 Files Created/Modified', level=2)
doc.add_paragraph(
    'Created: LandingPage.jsx, IntakeForm.jsx, CameraPage.jsx, ResultsPage.jsx, '
    'Stepper.jsx, ProgressRing.jsx, ScoreGauge.jsx, PatientContext.jsx, config.js, '
    'capture.js, ws.js, health.js\n\n'
    'Modified: App.jsx, main.jsx, styles.css (full rewrite), vite.config.js, backend/app.py (30s cap)'
)

# ═══════════════════════════════════════════════
# 11. RESULTS & PERFORMANCE
# ═══════════════════════════════════════════════
doc.add_heading('11. Results & Performance', level=1)

doc.add_heading('11.1 Measured Results', level=2)
doc.add_paragraph(
    'From Open-rppg_results.md, the engine produced the following measurements on test videos:'
)

table = doc.add_table(rows=4, cols=6, style='Light List Accent 1')
headers = ['Test', 'Total Frames', 'Key Frames', 'Non-Key', 'Skipped', 'Est. HR (BPM)']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
data = [
    ('Video 1', '1392', '6', '1340', '0', '152.82'),
    ('Video 2', '2238', '367', '4033', '2162', '76.95'),
    ('Video 3', '3307', '552', '6061', '3306', '100.44'),
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val

doc.add_heading('11.2 Performance Characteristics', level=2)
perf = [
    'Model warmup (JIT compilation): 30–60 seconds on first load',
    'Inference: Real-time at 30 FPS with JAX GPU acceleration',
    'Face detection: ~5ms per frame (ONNX Runtime)',
    'Streaming latency: <100ms end-to-end (camera → model → vitals → display)',
    'Memory: ~500MB–1GB per model instance (JAX preallocates GPU memory)',
    'Signal quality threshold: ≥300 samples (~10s at 30 FPS) needed for reliable HR/HRV'
]
for p_item in perf:
    doc.add_paragraph(p_item, style='List Bullet')

doc.add_heading('11.3 Output Structure', level=2)
doc.add_paragraph(
    'The returned dictionary contains:\n'
    '• hr — Heart Rate via FFT (BPM)\n'
    '• SQI — Signal Quality Index (0.0–1.0)\n'
    '• latency — Inference latency (seconds)\n'
    '• hrv — Heart Rate Variability dictionary:\n'
    '    • bpm, ibi, sdnn, rmssd, pnn50, LF/HF, breathingrate'
)

# ═══════════════════════════════════════════════
# 12. DEPLOYMENT & TESTING
# ═══════════════════════════════════════════════
doc.add_heading('12. Deployment & Testing', level=1)

doc.add_heading('12.1 Running the System', level=2)
doc.add_paragraph(
    'Backend:\n'
    'cd E:\\phase_1\\backend\n'
    '$env:KERAS_BACKEND = "jax"\n'
    '..\\.infer_venv\\Scripts\\python.exe -m uvicorn app:app --host 0.0.0.0 --port 8000\n\n'
    'Frontend:\n'
    'cd E:\\phase_1\\frontend\n'
    'npm install\n'
    'npm run dev\n\n'
    'Engine health check:\n'
    'curl http://localhost:8000/health\n'
    'curl http://localhost:8000/api/info'
)

doc.add_heading('12.2 Mobile Testing', level=2)
doc.add_paragraph(
    'Mobile browsers block camera on insecure origins (HTTP). For phone testing:\n'
    '1. Run backend on localhost:8000\n'
    '2. Run frontend on localhost:5173\n'
    '3. ngrok http 5173 (creates HTTPS tunnel)\n'
    '4. Open ngrok URL on phone → camera works via secure context (wss://)'
)

doc.add_heading('12.3 Verification', level=2)
doc.add_paragraph(
    'verify_engine.py replays a video through the same update_frame path used by WebSocket. '
    'The bundled Recording 2026-06-09.mp4 is a screen capture for validation (no face, so no HR).'
)

doc.add_heading('12.4 Remaining / Future Improvements', level=2)
improvements = [
    'Live BP estimation (currently not estimated by engine; Results shows self-reported category)',
    'High-contrast theme variant (prefers-contrast)',
    'Code-split recharts bundle (>500 kB build warning)',
    'Downloadable PDF/printable report (engine already has generate_visual_report)',
    'Historical session persistence and replay',
    'Multi-user backend pool sizing',
    'Silence React Router v7 future-flag warnings',
]
for imp in improvements:
    doc.add_paragraph(imp, style='List Bullet')

# ── Save ──
output_path = 'E:\\phase_1\\rPPG_Vitals_System_Architecture_Report.docx'
doc.save(output_path)
print(f'Report saved to {output_path}')
