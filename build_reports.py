from __future__ import annotations

import ast
import json
import math
import os
import re
import sys
import zipfile
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
_BUNDLED_SITE = Path.home() / ".cache/codex-runtimes/codex-primary-runtime/dependencies/python/Lib/site-packages"
if _BUNDLED_SITE.exists():
    sys.path.append(str(_BUNDLED_SITE))
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

try:
    import joblib
except Exception:
    joblib = None

try:
    import torch
except Exception:
    torch = None


ROOT = Path(__file__).resolve().parent
FIG = ROOT / "report_assets"
FIG.mkdir(exist_ok=True)


def safe_read(path: Path, limit=8000) -> str:
    try:
        return path.read_text(encoding="utf-8", errors="replace")[:limit]
    except Exception:
        return "Not Found in Repository"


def csv_info(path: Path) -> dict:
    try:
        df = pd.read_csv(path)
    except Exception as exc:
        return {"path": str(path.relative_to(ROOT)), "error": str(exc)}
    return {
        "path": str(path.relative_to(ROOT)),
        "rows": len(df),
        "columns": list(df.columns),
        "dtypes": {c: str(df[c].dtype) for c in df.columns},
        "missing": {c: int(df[c].isna().sum()) for c in df.columns},
        "df": df,
    }


def describe_column(name: str) -> str:
    n = name.lower()
    if n in {"subject_id", "trial_id"}:
        return "Identifier used for grouped splitting and traceability."
    if "hr" in n and "hrv" not in n:
        return "Heart-rate-derived feature or target."
    if "bvp" in n:
        return "Blood volume pulse waveform feature from Empatica BVP."
    if "ibi" in n or n in {"rmssd", "sdnn", "pnn50"}:
        return "Inter-beat interval / HRV feature."
    if "eda" in n:
        return "Electrodermal activity feature related to sympathetic arousal."
    if "temp" in n:
        return "Peripheral temperature feature related to perfusion and optical stability."
    if "acc" in n or "motion" in n:
        return "Accelerometer or motion-derived artifact/context feature."
    if "cv_load" in n:
        return "Pseudo cardiovascular-load target/index."
    if "bp_proxy" in n:
        return "Synthetic blood-pressure proxy target derived from load, motion, and reliability."
    if "sbp" in n:
        return "Systolic blood pressure value or prediction."
    if "dbp" in n:
        return "Diastolic blood pressure value or prediction."
    if "pred" in n:
        return "Model prediction output."
    if "actual" in n or "target" in n:
        return "Ground-truth or constructed target value."
    return "Repository column; semantic description not explicitly documented."


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_doc(doc: Document):
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    for style_name, size, color in [("Heading 1", 16, "2E74B5"), ("Heading 2", 13, "2E74B5"), ("Heading 3", 12, "1F4D78")]:
        s = doc.styles[style_name]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(10)
        s.paragraph_format.space_after = Pt(5)


def add_title(doc, title, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor.from_string("0B2545")
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(subtitle)
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor.from_string("555555")
    doc.add_paragraph()


def add_table(doc, headers, rows, widths=None, font_size=8):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = str(h)
        set_cell_shading(hdr[i], "E8EEF5")
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(font_size)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = "" if val is None else str(val)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(font_size)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(str(item), style="List Bullet")


def save_plot(fig, name):
    path = FIG / name
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)
    return path


def numeric_corr_plot(df, name, title):
    num = df.select_dtypes(include=np.number)
    if num.shape[1] < 2:
        return None
    corr = num.corr(numeric_only=True).fillna(0)
    fig, ax = plt.subplots(figsize=(8, 6))
    im = ax.imshow(corr, cmap="coolwarm", vmin=-1, vmax=1)
    ax.set_xticks(range(len(corr.columns)))
    ax.set_xticklabels(corr.columns, rotation=90, fontsize=6)
    ax.set_yticks(range(len(corr.columns)))
    ax.set_yticklabels(corr.columns, fontsize=6)
    ax.set_title(title)
    fig.colorbar(im, ax=ax, fraction=0.046, pad=0.04)
    return save_plot(fig, name)


def hist_plot(df, cols, name, title):
    cols = [c for c in cols if c in df.columns]
    if not cols:
        return None
    fig, axes = plt.subplots(1, len(cols), figsize=(4 * len(cols), 3))
    if len(cols) == 1:
        axes = [axes]
    for ax, c in zip(axes, cols):
        pd.to_numeric(df[c], errors="coerce").dropna().hist(ax=ax, bins=min(12, max(3, int(np.sqrt(len(df))))))
        ax.set_title(c)
        ax.set_xlabel("Value")
        ax.set_ylabel("Count")
    fig.suptitle(title)
    return save_plot(fig, name)


def pred_plot(df, actual, pred, name, title):
    if actual not in df or pred not in df:
        return None
    x = pd.to_numeric(df[actual], errors="coerce")
    y = pd.to_numeric(df[pred], errors="coerce")
    m = x.notna() & y.notna()
    if not m.any():
        return None
    fig, ax = plt.subplots(figsize=(5, 4))
    ax.scatter(x[m], y[m], s=28, alpha=0.8)
    lo, hi = min(x[m].min(), y[m].min()), max(x[m].max(), y[m].max())
    ax.plot([lo, hi], [lo, hi], color="gray", linestyle="--")
    ax.set_xlabel(actual)
    ax.set_ylabel(pred)
    ax.set_title(title)
    return save_plot(fig, name)


def error_metrics(df, actual, pred):
    if actual not in df or pred not in df:
        return None
    a = pd.to_numeric(df[actual], errors="coerce")
    p = pd.to_numeric(df[pred], errors="coerce")
    m = a.notna() & p.notna()
    if not m.any():
        return None
    err = p[m] - a[m]
    return {
        "n": int(m.sum()),
        "MAE": float(np.mean(np.abs(err))),
        "RMSE": float(np.sqrt(np.mean(err**2))),
        "R2": float(1 - np.sum(err**2) / np.sum((a[m] - a[m].mean()) ** 2)) if m.sum() > 1 and np.sum((a[m] - a[m].mean()) ** 2) else np.nan,
    }


def notebook_text(path):
    try:
        nb = json.loads(path.read_text(encoding="utf-8", errors="replace"))
    except Exception:
        return ""
    parts = []
    for cell in nb.get("cells", []):
        src = "".join(cell.get("source", []))
        if cell.get("cell_type") == "markdown":
            parts.append(src[:1000])
        for out in cell.get("outputs", []):
            if "text" in out:
                parts.append("".join(out["text"])[:1000])
    return "\n".join(parts)


def parse_open_rppg_notebook(path: Path) -> list[dict]:
    try:
        nb = json.loads(path.read_text(encoding="utf-8", errors="replace"))
    except Exception:
        return []
    cases = []
    for cell in nb.get("cells", []):
        src = "".join(cell.get("source", []))
        if "process_video" not in src:
            continue
        video_match = re.search(r'process_video\(["\']([^"\']+)["\']\)', src)
        outputs = "\n".join("".join(out.get("text", [])) for out in cell.get("outputs", []) if "text" in out)
        hr_match = re.search(r"Estimated Heart Rate:\s*([0-9.]+)\s*BPM", outputs)
        stat = {}
        for key in ["Total Frames", "Key Frames", "Non-Key Frames", "Skipped Frames", "Forward Filled Frames", "No Face Detected Frames"]:
            m = re.search(rf"{re.escape(key)}:\s*([0-9]+)", outputs)
            if m:
                stat[key] = int(m.group(1))
        warnings = []
        if "Frame rate mismatch" in outputs:
            warnings.append("Frame rate mismatch")
        if "non-key frames" in outputs:
            warnings.append("Detected non-key frames")
        cases.append(
            {
                "video": Path(video_match.group(1)).name if video_match else "Not Found in Repository",
                "path": video_match.group(1) if video_match else "Not Found in Repository",
                "hr": float(hr_match.group(1)) if hr_match else None,
                "stats": stat,
                "warnings": "; ".join(warnings) if warnings else "None recorded",
            }
        )
    return cases


def parse_open_rppg_markdown(path: Path) -> list[dict]:
    text = safe_read(path, 30000)
    blocks = re.split(r"Estimated Heart Rate:\s*([0-9.]+)\s*BPM", text)
    cases = []
    for i in range(1, len(blocks), 2):
        context = blocks[i - 1]
        hr = float(blocks[i])
        stat = {}
        for key in ["Total Frames", "Key Frames", "Non-Key Frames", "Skipped Frames", "Forward Filled Frames", "No Face Detected Frames"]:
            matches = re.findall(rf"{re.escape(key)}:\s*([0-9]+)", context)
            if matches:
                stat[key] = int(matches[-1])
        cases.append({"case": len(cases) + 1, "hr": hr, "stats": stat})
    return cases


def extract_classes(path):
    try:
        tree = ast.parse(path.read_text(encoding="utf-8", errors="replace"))
    except Exception:
        return []
    out = []
    for node in ast.walk(tree):
        if isinstance(node, ast.ClassDef):
            bases = [getattr(b, "id", getattr(getattr(b, "attr", None), "id", "")) for b in node.bases]
            out.append((node.name, ", ".join([b for b in bases if b]) or "object"))
    return out


def maybe_load_checkpoint(path):
    if torch is None:
        return {"error": "torch not available in document build runtime"}
    try:
        return torch.load(path, map_location="cpu", weights_only=False)
    except Exception as exc:
        return {"error": str(exc)}


def build_custom():
    features = csv_info(ROOT / "outputs/features/trial_features.csv")
    dl_windows = csv_info(ROOT / "outputs/features/dl_windows.csv")
    ml_pred = csv_info(ROOT / "outputs/predictions/ml_random_sample_predictions.csv")
    dl_pred = csv_info(ROOT / "outputs/predictions/dl_random_sample_predictions.csv")
    csvs = [features, dl_windows, ml_pred, dl_pred]
    source_files = [ROOT / "src/models/ml_models.py", ROOT / "src/models/dl_models.py", ROOT / "src/models/dl_training.py", ROOT / "src/models/datasets.py", ROOT / "src/features/build_features.py", ROOT / "src/features/pseudo_targets.py"]
    ckpt = maybe_load_checkpoint(ROOT / "outputs/models/multimodal_physio_net.pt")

    fdf = features.get("df", pd.DataFrame())
    figs = []
    if not fdf.empty:
        figs.append(numeric_corr_plot(fdf, "custom_correlation.png", "Trial Feature Correlation Matrix"))
        figs.append(hist_plot(fdf, ["hr_mean", "cv_load_index", "bp_proxy_score"], "custom_targets.png", "Custom Target / Proxy Distributions"))
        # correlation ranking
        target = "bp_proxy_score" if "bp_proxy_score" in fdf else None
        ranking = []
        if target:
            corr = fdf.select_dtypes(include=np.number).corr(numeric_only=True)[target].dropna().abs().sort_values(ascending=False)
            ranking = [(k, f"{v:.3f}") for k, v in corr.items() if k != target][:15]
        else:
            ranking = []
    else:
        ranking = []
    if "df" in ml_pred:
        figs.append(pred_plot(ml_pred["df"], "actual_hr", "predicted_hr", "ml_hr_pred.png", "ML HR: Actual vs Predicted"))
        figs.append(pred_plot(ml_pred["df"], "actual_bp_proxy", "predicted_bp_proxy", "ml_bp_pred.png", "ML BP Proxy: Actual vs Predicted"))
    if "df" in dl_pred:
        figs.append(pred_plot(dl_pred["df"], "actual_hr", "predicted_hr", "dl_hr_pred.png", "DL HR: Actual vs Predicted"))
        figs.append(pred_plot(dl_pred["df"], "actual_bp_proxy", "predicted_bp_proxy", "dl_bp_pred.png", "DL BP Proxy: Actual vs Predicted"))
    figs = [f for f in figs if f]

    doc = Document()
    style_doc(doc)
    add_title(doc, "Custom BP Estimation Analysis Report", "Research and technical audit of ML and DL pipelines")
    doc.add_heading("Executive Summary", level=1)
    add_bullets(doc, [
        "The repository implements a multimodal physiology pipeline over Empatica E4 signals and archived videos.",
        "The custom target is a BP proxy score, not cuff-measured SBP/DBP; direct BP ground truth for the custom pipeline was Not Found in Repository.",
        "Classical ML uses grouped subject-aware splitting with imputation, standardization, and XGBoost where available, otherwise Random Forest fallback.",
        "The custom DL model is a PyTorch multimodal network with video, BVP, physiology, static-feature, temporal-fusion, HR, cardiovascular-load, and BP-proxy heads.",
    ])
    doc.add_heading("Repository Evidence Base", level=1)
    add_table(doc, ["Artifact", "Status"], [
        ["Source modules", ", ".join(str(p.relative_to(ROOT)) for p in source_files if p.exists())],
        ["Notebooks", f"{len(list((ROOT/'notebooks').glob('*.ipynb')))} notebooks found"],
        ["Saved models", ", ".join(p.name for p in (ROOT/'outputs/models').glob('*'))],
        ["Prediction outputs", ", ".join(p.name for p in (ROOT/'outputs/predictions').glob('*.csv'))],
    ], [2.0, 4.2])

    doc.add_heading("Dataset Analysis", level=1)
    rows = []
    for info in csvs:
        rows.append([info["path"], info.get("rows", "Not Found"), len(info.get("columns", [])), ", ".join(info.get("columns", [])[:8])])
    add_table(doc, ["CSV", "Rows", "Columns", "Column sample"], rows, [2.1, .7, .8, 2.9])
    if not fdf.empty:
        subjects = fdf["subject_id"].nunique() if "subject_id" in fdf else "Not Found in Repository"
        trials = len(fdf)
        doc.add_paragraph(f"Trial-level feature table contains {trials} rows and {subjects} unique subjects. Grouped train/test splitting is implemented in source code with GroupShuffleSplit(test_size=0.3, random_state=42).")

    doc.add_heading("CSV Column Dictionary", level=2)
    for info in csvs:
        doc.add_heading(info["path"], level=3)
        rows = [[c, info.get("dtypes", {}).get(c, ""), describe_column(c), "Feature/target/prediction as implied by name and code"] for c in info.get("columns", [])[:40]]
        add_table(doc, ["Column Name", "Data Type", "Description", "Usage"], rows, [1.5, 1.0, 2.2, 1.7], 7)

    doc.add_heading("Exploratory Data Analysis", level=1)
    if not fdf.empty:
        miss = fdf.isna().sum().sort_values(ascending=False)
        add_table(doc, ["Column", "Missing Count", "Missing %"], [(c, int(v), f"{100*v/len(fdf):.1f}%") for c, v in miss.items() if v > 0][:20] or [["All inspected columns", 0, "0.0%"]], [2.5, 1.3, 1.2])
        num = fdf.select_dtypes(include=np.number)
        out_rows = []
        for c in num.columns:
            q1, q3 = num[c].quantile(.25), num[c].quantile(.75)
            iqr = q3 - q1
            if iqr and np.isfinite(iqr):
                out = ((num[c] < q1 - 1.5 * iqr) | (num[c] > q3 + 1.5 * iqr)).sum()
                if out:
                    out_rows.append([c, int(out), "IQR 1.5x rule"])
        add_table(doc, ["Feature", "Outlier Count", "Method"], out_rows[:20] or [["No numeric outliers flagged", 0, "IQR 1.5x rule"]], [2.4, 1.2, 2.0])
        if ranking:
            add_table(doc, ["Feature", "Absolute correlation with BP proxy"], ranking, [3.5, 2.0])
    for fig in figs:
        doc.add_picture(str(fig), width=Inches(6.1))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Feature Engineering Analysis", level=1)
    add_table(doc, ["Feature Family", "Actual Features Found", "Meaning"], [
        ["Signal summaries", "hr_mean/std/min/max, eda_mean/std/min/max, temp_mean/std/min/max", "Central tendency and variability across trial signals."],
        ["Motion", "motion_mean/std/min/max", "Magnitude of ACC channels used as motion context and artifact proxy."],
        ["HRV", "rmssd, sdnn, pnn50 and related IBI-derived outputs when computable", "Autonomic regulation and beat-to-beat variability."],
        ["BVP", "bvp_* plus bvp_snr", "Pulse morphology and signal quality from contact PPG."],
        ["Pseudo targets", "cv_load_index/class, signal_reliability_score, stress_recovery_index, bp_proxy_score, hr_baseline_deviation", "Constructed labels for first-phase modeling."],
    ], [1.5, 2.5, 2.2])

    doc.add_heading("Machine Learning Model Analysis", level=1)
    add_table(doc, ["Component", "Repository Evidence"], [
        ["Algorithms", "XGBRegressor/XGBClassifier with n_estimators=100, max_depth=3, learning_rate=0.05, subsample=0.9; RandomForest fallback with n_estimators=200."],
        ["Preprocessing", "SimpleImputer(strategy='median') and StandardScaler in sklearn Pipeline."],
        ["Splitting", "GroupShuffleSplit by subject_id when possible, test_size=0.3, random_state=42."],
        ["Targets", "HR regression, cardiovascular-load classification, BP proxy regression."],
    ], [1.8, 4.4])
    for model_file in ["hr_regressor.joblib", "cv_load_classifier.joblib", "bp_proxy_regressor.joblib"]:
        path = ROOT / "outputs/models" / model_file
        if path.exists():
            if joblib is None:
                doc.add_paragraph(f"{model_file}: saved artifact found; joblib not available in document build runtime, so object repr was not loaded.")
            else:
                try:
                    model = joblib.load(path)
                    doc.add_paragraph(f"{model_file}: {model}")
                except Exception:
                    doc.add_paragraph(f"{model_file}: saved artifact found; detailed load failed.")

    doc.add_heading("Deep Learning Model Analysis", level=1)
    classes = extract_classes(ROOT / "src/models/dl_models.py")
    add_table(doc, ["Class", "Base"], classes, [2.5, 2.5])
    cfg = ckpt.get("config", {}) if isinstance(ckpt, dict) else {}
    hist = ckpt.get("history", []) if isinstance(ckpt, dict) else []
    add_table(doc, ["Training/Checkpoint Item", "Value"], [
        ["Checkpoint config", json.dumps(cfg) if cfg else "Not Found in Repository"],
        ["History rows", len(hist) if hist else "Not Found in Repository"],
        ["Optimizer/loss in code", "AdamW lr default 1e-3, weight_decay=1e-4; SmoothL1 HR/BP losses and CrossEntropy cardiovascular-load loss."],
        ["Default epochs", "fit_multitask_model default epochs=5; notebook may override if present."],
    ], [2.3, 3.8])
    if hist:
        hdf = pd.DataFrame(hist)
        add_table(doc, list(hdf.columns), hdf.round(4).astype(str).values.tolist()[:20], font_size=7)
        if "epoch" in hdf and "train_loss" in hdf:
            fig, ax = plt.subplots(figsize=(6, 3.5))
            ax.plot(hdf["epoch"], hdf["train_loss"], label="train_loss")
            if "val_loss" in hdf:
                ax.plot(hdf["epoch"], hdf["val_loss"], label="val_loss")
            ax.set_title("DL Training Curve")
            ax.set_xlabel("Epoch")
            ax.set_ylabel("Loss")
            ax.legend()
            path = save_plot(fig, "dl_training_curve.png")
            doc.add_picture(str(path), width=Inches(6.1))

    doc.add_heading("Prediction Result Analysis", level=1)
    rows = []
    for label, info in [("ML HR", ml_pred), ("ML BP proxy", ml_pred), ("DL HR", dl_pred), ("DL BP proxy", dl_pred)]:
        df = info.get("df", pd.DataFrame())
        if "HR" in label:
            m = error_metrics(df, "actual_hr", "predicted_hr")
        else:
            m = error_metrics(df, "actual_bp_proxy", "predicted_bp_proxy")
        if m:
            rows.append([label, m["n"], f'{m["MAE"]:.4f}', f'{m["RMSE"]:.4f}', f'{m["R2"]:.4f}'])
        else:
            rows.append([label, "Not Found", "Not Found", "Not Found", "Not Found"])
    add_table(doc, ["Model/Target", "N", "MAE", "RMSE", "R2"], rows, [1.8, .6, 1.1, 1.1, 1.1])

    doc.add_heading("Comparative Analysis", level=1)
    doc.add_paragraph("ML is more transparent and less computationally demanding on the current small tabular feature set. DL is architecturally richer and ready for video/physiology fusion, but its generalization is limited by the small number of available subjects and the absence of direct cuff BP targets in the custom pipeline.")
    doc.add_heading("Discussion and Limitations", level=1)
    add_bullets(doc, [
        "Direct SBP/DBP custom-model labels were Not Found in Repository; the custom task estimates a BP proxy.",
        "The archive contains a small subject/trial count, so model metrics should be treated as first-pass engineering validation.",
        "Feature definitions are traceable in code, but several notebook-level results are not persisted as formal experiment logs.",
    ])
    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph("The best-supported conclusion is that the repository successfully builds a modular first-phase multimodal physiology pipeline. The classical ML pipeline is the most auditable current model; the DL model is a credible experimental architecture but requires more subjects and real BP labels for publication-strength BP estimation claims.")
    out = ROOT / "Custom_BP_Estimation_Analysis_Report.docx"
    doc.save(out)
    return out


def build_pretrained():
    vbpe_csvs = [csv_info(p) for p in [ROOT/"rppg_inference_vbpe/SBP_new.csv", ROOT/"rppg_inference_vbpe/DBP_new.csv", ROOT/"rppg_inference_vbpe/Input_Data/Demographic_Data.csv"] if p.exists()]
    ppg_csvs = [csv_info(p) for p in sorted((ROOT/"rppg_inference_vbpe/Results/PPG").glob("**/*.csv"))]
    open_results = safe_read(ROOT / "open_rppg_inference/Open-rppg_results.md", 10000)
    vbpe_results = safe_read(ROOT / "rppg_inference_vbpe/Results/Results.md", 10000)
    open_cases = parse_open_rppg_notebook(ROOT / "open_rppg_inference/open_rppg_inference.ipynb")
    open_md_cases = parse_open_rppg_markdown(ROOT / "open_rppg_inference/Open-rppg_results.md")

    figs = []
    for info in vbpe_csvs:
        df = info.get("df", pd.DataFrame())
        if not df.empty:
            figs.append(hist_plot(df, df.select_dtypes(include=np.number).columns[:4].tolist(), f"{Path(info['path']).stem}_dist.png", f"{info['path']} Numeric Distributions"))

    doc = Document()
    style_doc(doc)
    add_title(doc, "Pretrained Model Analysis Report", "Evaluation and comparison of V-BPE and Open-rPPG inference systems")
    doc.add_heading("Executive Summary", level=1)
    add_bullets(doc, [
        "rppg_inference_vbpe is a video-based blood pressure estimation pipeline using MediaPipe face/hand landmarks, DeepPhys-style PPG extraction, demographic/vessel-length terms, and SBP/DBP calculation.",
        "open_rppg_inference is a general rPPG toolbox for HR/HRV and BVP extraction with a model zoo including FacePhys, PhysMamba, RhythmMamba, PhysFormer, TSCAN, PhysNet, and EfficientPhys.",
        "Only V-BPE contains explicit BP prediction outputs. Open-rPPG BP estimation outputs were Not Found in Repository; it should be treated as a physiological-signal extractor unless paired with a BP model.",
    ])
    doc.add_heading("Codebase Inspection", level=1)
    add_table(doc, ["Folder", "Framework / Dependencies", "Model Type", "Primary Output"], [
        ["rppg_inference_vbpe", "Python, OpenCV, MediaPipe, PyTorch checkpoint can2d_pytorch.pth", "Face/hand PTT + PPG-based analytical BP estimator", "SBP_new.csv, DBP_new.csv"],
        ["open_rppg_inference", "Python, JAX-style model loaders, ONNXRuntime face detector, PyAV, HeartPy, SciPy", "rPPG deep model zoo and BVP/HR/HRV extraction", "HR, SQI, HRV, BVP; BP Not Found in Repository"],
    ], [1.5, 2.1, 1.8, 1.5])
    doc.add_heading("Open-RPPG Notebook and Result File Inspection", level=2)
    doc.add_paragraph("The notebook open_rppg_inference.ipynb installs open-rppg, initializes the default rppg.Model(), and runs offline process_video inference on the same sample identifiers used by the V-BPE folder. The default model maps to FacePhys.rlap in open_rppg_inference/rppg/main.py, which loads FacePhys.rlap.weights.h5 and uses a 30 FPS, 36 x 36 face input.")
    if open_cases:
        add_table(doc, ["Sample Video", "Notebook Path", "Estimated HR (BPM)", "Total Frames", "Key Frames", "Non-Key Frames", "Skipped", "Forward Filled", "No Face", "Warnings"], [
            [
                c["video"],
                c["path"],
                f"{c['hr']:.2f}" if c["hr"] is not None else "Not Found",
                c["stats"].get("Total Frames", "Not Found"),
                c["stats"].get("Key Frames", "Not Found"),
                c["stats"].get("Non-Key Frames", "Not Found"),
                c["stats"].get("Skipped Frames", "Not Found"),
                c["stats"].get("Forward Filled Frames", "Not Found"),
                c["stats"].get("No Face Detected Frames", "Not Found"),
                c["warnings"],
            ]
            for c in open_cases
        ], [.8, 1.15, .75, .55, .55, .65, .55, .65, .55, 1.1], 6)
    if open_md_cases:
        add_table(doc, ["Markdown Case", "Estimated HR (BPM)", "Total Frames", "Key Frames", "Non-Key Frames", "Skipped", "Forward Filled", "No Face"], [
            [
                c["case"],
                f"{c['hr']:.2f}",
                c["stats"].get("Total Frames", "Not Found"),
                c["stats"].get("Key Frames", "Not Found"),
                c["stats"].get("Non-Key Frames", "Not Found"),
                c["stats"].get("Skipped Frames", "Not Found"),
                c["stats"].get("Forward Filled Frames", "Not Found"),
                c["stats"].get("No Face Detected Frames", "Not Found"),
            ]
            for c in open_md_cases
        ], [.8, 1.0, .8, .75, .85, .7, .8, .7], 7)
    if open_cases:
        hrs = [c["hr"] for c in open_cases if c["hr"] is not None]
        if hrs:
            fig, ax = plt.subplots(figsize=(6, 3.4))
            ax.bar([c["video"] for c in open_cases if c["hr"] is not None], hrs, color="#2E74B5")
            ax.set_ylabel("Estimated HR (BPM)")
            ax.set_title("Open-RPPG FacePhys.rlap HR Estimates by Sample")
            ax.tick_params(axis="x", rotation=25)
            hr_fig = save_plot(fig, "open_rppg_hr_by_sample.png")
            doc.add_picture(str(hr_fig), width=Inches(6.1))
        warning_count = sum(1 for c in open_cases if c["warnings"] != "None recorded")
        doc.add_paragraph(f"Interpretation: Open-RPPG detected faces in all notebook cases because No Face Detected Frames is 0 for each logged run. However, {warning_count} of {len(open_cases)} notebook runs include video-quality warnings. The SL377 case reports heavy non-key-frame use and forward filling; SL421, SL428, and SL413 additionally report frame-rate mismatch and large skipped-frame counts. These warnings matter because Open-RPPG estimates HR from temporal BVP frequency content, so irregular sampling and inter-frame compression can distort the pulse waveform.")
    else:
        doc.add_paragraph("Open-RPPG notebook case extraction: Not Found in Repository.")
    doc.add_heading("Architecture and Pipeline Diagrams", level=1)
    add_table(doc, ["System", "Workflow"], [
        ["V-BPE", "Input video -> MediaPipe face/hand landmarking -> cropped face/hand videos -> MP4 conversion -> DeepPhys PPG extraction -> face/hand PTT -> vessel/demographic terms -> SBP and DBP CSV outputs."],
        ["Open-rPPG", "Input video/tensor -> ONNX face detection -> face crop buffering/resampling -> selected pretrained model inference -> BVP filtering/normalization -> HR, SQI, and HRV metrics."],
    ], [1.3, 5.0])
    doc.add_heading("Dataset Analysis", level=1)
    add_table(doc, ["Data Source", "Count / Status"], [
        ["V-BPE input videos", len(list((ROOT/"rppg_inference_vbpe/Input_Videos").glob("*.mp4")))],
        ["V-BPE input photos", len(list((ROOT/"rppg_inference_vbpe/Input_Photos").glob("*.*")))],
        ["V-BPE PPG CSVs", len(ppg_csvs)],
        ["Open-RPPG notebook inference cases", len(open_cases) if open_cases else "Not Found in Repository"],
        ["Open-rPPG direct BP labels", "Not Found in Repository"],
    ], [2.5, 2.0])
    for info in vbpe_csvs + ppg_csvs[:4]:
        doc.add_heading(info["path"], level=2)
        add_table(doc, ["Column Name", "Data Type", "Description", "Usage"], [[c, info.get("dtypes", {}).get(c, ""), describe_column(c), "Pretrained pipeline data/result field"] for c in info.get("columns", [])[:25]], [1.5, 1, 2.3, 1.5], 7)
    for fig in [f for f in figs if f]:
        doc.add_picture(str(fig), width=Inches(6.1))

    doc.add_heading("rPPG Pipeline Analysis", level=1)
    add_table(doc, ["Stage", "V-BPE", "Open-rPPG"], [
        ["Face Detection / Tracking", "MediaPipe face landmarker; separate hand capture.", "ONNX BlazeFace detector with Kalman-smoothed boxes."],
        ["ROI Selection", "Face and hand cropped regions saved as videos.", "Detected face crop resized to model input resolution."],
        ["Signal Extraction", "can2dshare.py invokes DeepPhys/CAN-like PPG extraction.", "Selected pretrained rPPG model emits BVP-like signal."],
        ["Filtering", "PPG CSV generation and downstream BP scripts; details partly in source.", "Bandpass 0.5-3 Hz, detrending, normalization, SQI."],
        ["BP Estimation", "SBP and DBP scripts present and CSV outputs saved.", "Not Found in Repository."],
    ], [1.3, 2.45, 2.45], 8)
    doc.add_heading("Open-RPPG Inference Workflow Details", level=2)
    add_table(doc, ["Step", "Code / Notebook Evidence", "Technical Interpretation"], [
        ["Model initialization", "model = rppg.Model() in every notebook sample; main.py default is FacePhys.rlap.", "All reported notebook outputs use the same default pretrained FacePhys.rlap model unless the notebook was edited outside the stored cells."],
        ["Video decoding", "Model.process_video opens the file with PyAV and iterates decoded frames.", "Input format can be mp4/mkv; timestamps are read from decoded frames."],
        ["Face localization", "FaceDetector uses ONNX BlazeFace and updates face boxes every 5 frames with Kalman smoothing.", "The output logs show 0 no-face frames in stored cases, so face detection did not fail in these runs."],
        ["Face resampling", "Detected face crops are resized to the model input size from metadata.", "For FacePhys.rlap, local code reports input=(1, 36, 36, 3), enabling low-latency frame-wise inference."],
        ["Signal processing", "bvp() applies optional cumulative handling, detrending, bandpass filtering, and normalization.", "The final HR is computed from the Welch frequency-domain peak in the 30-180 BPM band."],
        ["Quality checks", "process_video logs frame-rate mismatch, non-key frames, skipped frames, filled frames, and no-face frames.", "The stored results should be interpreted as HR estimates under imperfect video encoding conditions, not validated BP predictions."],
    ], [1.25, 2.45, 2.5], 7)
    doc.add_heading("Model Architecture Analysis", level=1)
    add_table(doc, ["Folder", "Architectures Found"], [
        ["rppg_inference_vbpe", "can2dshare.py and can2d_pytorch.pth indicate a CAN/DeepPhys-style PPG extractor; bp_calc.py/dbp.py implement BP formulas; MediaPipe task files support landmarking."],
        ["open_rppg_inference", "Supported model list includes FacePhys, ME chunk/flow, PhysMamba, RhythmMamba, PhysFormer, TSCAN, PhysNet, EfficientPhys. Local weights are present for pure/rlap variants."],
    ], [1.5, 4.7])
    doc.add_heading("Output Analysis", level=1)
    rows = []
    for info in vbpe_csvs:
        df = info.get("df", pd.DataFrame())
        nums = df.select_dtypes(include=np.number)
        if nums.empty:
            rows.append([info["path"], len(df), "No numeric columns"])
        else:
            desc = "; ".join(f"{c}: mean={nums[c].mean():.2f}, std={nums[c].std():.2f}" for c in nums.columns[:4])
            rows.append([info["path"], len(df), desc])
    add_table(doc, ["Output", "Rows", "Statistics"], rows or [["Not Found in Repository", "", ""]], [2.0, .7, 3.6], 8)
    if open_cases:
        hr_values = [c["hr"] for c in open_cases if c["hr"] is not None]
        add_table(doc, ["Open-RPPG HR Statistic", "Value"], [
            ["Samples analyzed in notebook", len(open_cases)],
            ["Mean estimated HR", f"{np.mean(hr_values):.2f} BPM" if hr_values else "Not Found"],
            ["Minimum estimated HR", f"{np.min(hr_values):.2f} BPM" if hr_values else "Not Found"],
            ["Maximum estimated HR", f"{np.max(hr_values):.2f} BPM" if hr_values else "Not Found"],
            ["Direct SBP/DBP output", "Not Found in Repository"],
            ["Ground-truth HR/BP comparison for Open-RPPG cases", "Not Found in Repository"],
        ], [2.5, 2.2])
    doc.add_paragraph("V-BPE results markdown excerpt: " + (re.sub(r"\s+", " ", vbpe_results)[:900] if vbpe_results else "Not Found in Repository"))
    doc.add_paragraph("Open-rPPG results markdown excerpt: " + (re.sub(r"\s+", " ", open_results)[:900] if open_results else "Not Found in Repository"))

    doc.add_heading("Comparative Analysis", level=1)
    add_table(doc, ["Criterion", "V-BPE", "Open-rPPG"], [
        ["BP relevance", "Direct SBP/DBP pipeline and outputs.", "No BP estimator found; HR/HRV/BVP only."],
        ["Architecture breadth", "Focused face/hand PTT BP estimation.", "Broader model zoo with modern rPPG architectures."],
        ["Robustness controls", "Demographics/vessel length plus face/hand regions.", "SQI, face tracking, frame-rate checks, multiple pretrained models."],
        ["Inference complexity", "Multi-stage video preprocessing plus BP scripts.", "Unified API; complexity depends on selected model."],
        ["Best use in this repository", "BP prediction/evaluation.", "Reusable rPPG feature extractor and comparator."],
    ], [1.6, 2.3, 2.3], 8)
    doc.add_heading("Final Findings", level=1)
    add_bullets(doc, [
        "Better pretrained BP model in this repository: rppg_inference_vbpe, because it actually produces SBP/DBP outputs.",
        "Stronger general rPPG architecture suite: open_rppg_inference, because it packages multiple modern pretrained rPPG networks.",
        "Best combined path: use Open-rPPG for robust BVP/quality features and V-BPE-style calibration/ground-truth BP evaluation for BP-specific modeling.",
        "Missing information: formal runtime benchmarks, memory usage, and direct shared BP error metrics against cuff labels were Not Found in Repository.",
    ])
    out = ROOT / "Pretrained_Model_Analysis_Report.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    print(build_custom())
    print(build_pretrained())
